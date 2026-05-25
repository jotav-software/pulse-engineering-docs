#!/usr/bin/env python3
"""Gera docs/especificacao-funcional/*.md a partir do docx (sem strikethrough) e backup."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "especificacao-funcional"
DOCX = ROOT / "docs" / "especificacao_funcional.docx"
BACKUP = ROOT / "docs" / "especificacao_funcional_backup_before_format.docx"


def para_text(paragraph) -> str | None:
    parts = []
    for run in paragraph.runs:
        if run.font.strike:
            continue
        parts.append(run.text)
    t = "".join(parts).strip()
    return t or None


def cell_text(cell) -> str:
    parts = []
    for p in cell.paragraphs:
        t = para_text(p)
        if t:
            parts.append(t)
    return " ".join(parts).strip()


def table_md(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    hdr = rows[0]
    n = len(hdr)
    lines = [
        "| " + " | ".join(hdr) + " |",
        "| " + " | ".join(["---"] * n) + " |",
    ]
    for row in rows[1:]:
        cells = (row + [""] * n)[:n]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def load_blocks(path: Path) -> list[tuple[str, int, object]]:
    doc = Document(path)
    blocks: list[tuple[str, int, object]] = []
    pi, ti = 0, 0
    for child in doc.element.body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = doc.paragraphs[pi]
            pi += 1
            t = para_text(p)
            if not t:
                continue
            style = p.style.name if p.style else ""
            level = 0
            if style.startswith("Heading"):
                m = re.search(r"(\d+)", style)
                level = int(m.group(1)) if m else 1
            blocks.append(("p", level, t))
        elif tag == "tbl":
            table = doc.tables[ti]
            ti += 1
            rows = [[cell_text(c) for c in row.cells] for row in table.rows]
            blocks.append(("t", 0, rows))
    return blocks


def section_blocks(
    blocks: list[tuple[str, int, object]], start: str, end: str | None = None
) -> list[tuple[str, int, object]]:
    out: list[tuple[str, int, object]] = []
    cap = False
    for kind, level, data in blocks:
        if kind == "p" and level == 1 and isinstance(data, str):
            if start in data:
                cap = True
            elif cap and (end is None or end in data):
                if end and end in data:
                    break
                if end is None or start not in data:
                    break
        if cap:
            out.append((kind, level, data))
    return out


def hu_from_backup(hu_title: str) -> dict[str, list[str]]:
    doc = Document(BACKUP)
    sections: dict[str, list[str]] = {
        "telas": [],
        "campos": [],
        "regras": [],
        "fluxo": [],
        "excecoes": [],
        "aceite": [],
    }
    current = None
    capturing = False
    epic_re = re.compile(r"^\d+\.\s+Épico\s+\d+")
    for p in doc.paragraphs:
        t = para_text(p)
        if not t:
            continue
        if hu_title in t and "HU" in t:
            capturing = True
            continue
        if capturing:
            if t.startswith("HU ") and hu_title not in t:
                break
            if epic_re.match(t) or (t.startswith("Épico ") and "—" in t):
                break
            if t.startswith("Features e status"):
                break
            key_map = {
                "Telas envolvidas": "telas",
                "Campos / componentes principais": "campos",
                "Regras de negócio": "regras",
                "Fluxo principal": "fluxo",
                "Cenários de exceção / erro": "excecoes",
                "Critérios de aceite": "aceite",
            }
            if t in key_map:
                current = key_map[t]
                continue
            if current and not t.startswith("Status:") and not t.startswith("Objetivo:"):
                if t.startswith("Observação:"):
                    sections.setdefault("obs", []).append(t)
                else:
                    sections[current].append(t)
    return sections


def render_hu(hu_title: str, meta: dict[str, str], sections: dict[str, list[str]]) -> str:
    lines = [f"### {hu_title} — {meta.get('status', '[IMPLEMENTADO]')}", ""]
    for k in ("objetivo", "origem", "permissoes"):
        if meta.get(k):
            lines.append(f"**{k.capitalize()}:** {meta[k]}")
            lines.append("")
    for label, key in [
        ("Telas", "telas"),
        ("Campos / componentes", "campos"),
        ("Regras de negócio", "regras"),
        ("Fluxo principal", "fluxo"),
        ("Exceções / erros", "excecoes"),
        ("Critérios de aceite", "aceite"),
    ]:
        items = sections.get(key, [])
        if items:
            lines.append(f"**{label}**")
            lines.append("")
            for i in items:
                lines.append(f"- {i}")
            lines.append("")
    if sections.get("obs"):
        for o in sections["obs"]:
            lines.append(f"> {o}")
            lines.append("")
    return "\n".join(lines)


def parse_hu_meta(blocks: list) -> dict[str, dict[str, str]]:
    metas: dict[str, dict[str, str]] = {}
    current_hu = None
    for kind, level, data in blocks:
        if kind != "p":
            continue
        if level == 2 and isinstance(data, str) and data.startswith("HU "):
            current_hu = data
            metas[current_hu] = {}
        elif current_hu and isinstance(data, str):
            if data.startswith("Status:"):
                metas[current_hu]["status"] = data.replace("Status:", "").strip()
            elif data.startswith("Objetivo:"):
                metas[current_hu]["objetivo"] = data.replace("Objetivo:", "").strip()
            elif data.startswith("Origem e acesso:"):
                metas[current_hu]["origem"] = data.replace("Origem e acesso:", "").strip()
            elif data.startswith("Permissões:"):
                metas[current_hu]["permissoes"] = data.replace("Permissões:", "").strip()
    return metas


def epic_tables(blocks: list) -> dict[str, str]:
    tables: dict[str, str] = {}
    last_h2 = ""
    for kind, level, data in blocks:
        if kind == "p" and level == 2:
            last_h2 = data
        if kind == "t" and last_h2.startswith("Features"):
            tables[last_h2] = table_md(data)  # type: ignore
    return tables


def write(path: Path, content: str) -> None:
    path.write_text(content.rstrip() + "\n", encoding="utf-8")


def main() -> None:
    blocks = load_blocks(DOCX)
    backup_blocks = load_blocks(BACKUP)
    OUT.mkdir(parents=True, exist_ok=True)

    # Shared legend
    legend = """| Tag | Significado |
| --- | --- |
| `[IMPLEMENTADO]` | Entregue e utilizável em produção ou demo estável |
| `[PARCIAL]` | Fluxo existe com lacunas (inclui UI «em breve») |
| `[PENDENTE]` | Não implementado ou apenas planejado |

Fonte de status: código (`app-producer`, `producer-web`, `app-client`, `client-web`, `backend`) + `docs/RBAC.md` + revisão 2026-05-19.
"""

    transversal = section_blocks(blocks, "3. Visão geral", "4. Registros")
    trans_lines = [t for k, _, t in transversal if k == "p" and _ == 0]

    # README
    write(
        OUT / "README.md",
        f"""# Especificação funcional — Ecossistema Pulse

> **Fonte de verdade** para regras de negócio, fluxos, permissões e backlog por plataforma.  
> Migrado de `docs/especificacao_funcional.docx` (conteúdo riscado ignorado). Arquivos `.docx` permanecem como arquivo histórico.

**Última revisão:** 2026-05-19

---

## Documentos por sistema

| Arquivo | Sistema | Público |
| --- | --- | --- |
| [pulse-admin.md](./pulse-admin.md) | Pulse Admin (backoffice) | Operadores `PULSE_ADMIN` |
| [app-produtor.md](./app-produtor.md) | App Produtor (mobile) | Dono, Gestor, Staff |
| [producer-web.md](./producer-web.md) | Producer Web (portal produtora) | Dono, Gestor, Staff |
| [app-client.md](./app-client.md) | App Cliente (mobile B2C) | Comprador, Promoter |
| [client-web.md](./client-web.md) | Client Web (site comprador) | Visitante / comprador |

**Transversal:** [arquitetura.md](./arquitetura.md) · [api-endpoints.md](./api-endpoints.md)

---

## Legenda de status

{legend}

---

## Glossário de plataformas

| Nome | Repo / rota | Observação |
| --- | --- | --- |
| **App Produtor** | `app-producer/` | Operação mobile (Access, financeiro completo) |
| **Producer Web** | `producer-web/` — `/dashboard`, `/events`, … | Portal web da produtora (sem `/admin`) |
| **Pulse Admin** | `producer-web/` — `/admin/*` | Backoffice interno; API `/api/admin/v1` |
| **App Cliente** | `app-client/` | Compra, carteira, facial, promoter |
| **Client Web** | `client-web/` | Vitrine pública; checkout [PENDENTE] |

---

## Regras transversais (todas as plataformas)

Consulte também `produto/regras-negocio/global-business-rules.md` (pulse-engineering-docs).

| Regra | Detalhe |
| --- | --- |
| Taxa Pulse | Repassada ao **comprador** nas vendas pela plataforma |
| Pix | 5% de desconto sobre a taxa da plataforma |
| Cartão | Até **4x**; juros de parcelamento repassados ao comprador |
| Limite compra | Máx. **4 ingressos por evento** por CPF/conta |
| Checkout | Reserva **10 min**; máx. **3 tentativas** de pagamento por pedido |
| Cancelamento comprador | Até **24h antes** do início; ingresso **USED** não reembolsa |
| Repasse produtor | Retido até elegibilidade (ver [payout-policies.md](../produto/regras-negocio/payout-policies.md)) |
| Facial vs QR | Facial principal; QR contingência obrigatória |
| Promoter | Experiência principal no **App Cliente** (`/promoter`) |

**Princípios invioláveis:** ingresso de venda só após `PAID`; emissão manual auditável; pedido expirado devolve estoque; `USED` não reverte; RBAC validado no backend.

---

## Mapa rápido de épicos (produtor)

| # | Épico | App Produtor | Producer Web |
| --- | --- | --- | --- |
| 1 | Acesso & Onboarding | [IMPLEMENTADO] | [IMPLEMENTADO] |
| 2 | Perfil & Configurações | [PARCIAL] | [PARCIAL] |
| 3 | Gestão de Eventos | [IMPLEMENTADO] | [IMPLEMENTADO] |
| 4 | Oferta Comercial | [IMPLEMENTADO] | [IMPLEMENTADO] |
| 5 | Lotes avançados | [IMPLEMENTADO] | [PARCIAL] |
| 6 | Dashboard | [IMPLEMENTADO] | [IMPLEMENTADO] |
| 7 | Pedidos & Participantes | [IMPLEMENTADO] | [IMPLEMENTADO] |
| 8 | Access / Check-in | [IMPLEMENTADO] | [PENDENTE] |
| 9 | Financeiro & Repasses | [IMPLEMENTADO] | [PARCIAL] |
| 10 | Equipe & RBAC | [PARCIAL] | [PARCIAL] |

Detalhes por sistema nos arquivos linkados acima.

---

## Referências canônicas

- [RBAC.md](../RBAC.md) — papéis e matriz por app
- [CHECKOUT_COMPLIANCE.md](../CHECKOUT_COMPLIANCE.md) — gate de termos B2C/produtor
- [global-business-rules.md](../produto/regras-negocio/global-business-rules.md) · [payout-policies.md](../produto/regras-negocio/payout-policies.md)
- OpenAPI: `GET /swagger` no backend
""",
    )

    # Architecture
    arch_blocks = section_blocks(blocks, "22. Arquitetura", "23. Recursos")
    arch_text = "\n".join(t for k, _, t in arch_blocks if k == "p" and _ == 0)
    write(
        OUT / "arquitetura.md",
        f"""# Arquitetura do ecossistema Pulse

> Escopo: visão técnica de alto nível | Público: engenharia e PO | Plataforma: monorepo | Última revisão: 2026-05-19

## Legenda de status

{legend}

## 1. Visão geral

Ecossistema **monorepo**: backend único (Elysia/Node + Prisma/MySQL), quatro frontends e serviço auxiliar de biometria.

```mermaid
flowchart LR
  subgraph clients [Compradores]
    AC[App Cliente Expo]
    CW[Client Web Next.js]
  end
  subgraph producers [Produtora]
    AP[App Produtor Expo]
    PW[Producer Web Next.js]
  end
  subgraph platform [Plataforma]
    API[Backend Elysia]
    DB[(MySQL)]
    PF[pulse-face]
  end
  subgraph external [Externos]
    PG[Pagar.me]
    BR[Brevo]
    BA[Better Auth]
  end
  AC --> API
  CW --> API
  AP --> API
  PW --> API
  API --> DB
  API --> PG
  API --> BR
  API --> BA
  API --> PF
```

## 2. Superfícies de API

| Prefixo | Consumidores | Descrição |
| --- | --- | --- |
| `/api/client/v1/*` | App Cliente, Client Web (futuro) | B2C canônico (OpenAPI) |
| Raiz legada (`/auth`, `/events`, …) | Apps antigos | Espelho; evitar em clientes novos |
| `/api/producer/v1/*` | App Produtor, Producer Web produtora | Portal + operação |
| `/api/admin/v1/*` | Pulse Admin | Restrito a `PULSE_ADMIN` |
| `/api/promoter/*` | App Cliente (promoter) | Comissões e vendas |
| `/api/auth/*` | Todos | Better Auth handler |
| `/internal/facial-*` | Jobs/cron | API key; galeria e retenção |

## 3. Autenticação e sessão

- **B2C:** Better Auth via `/api/client/v1/auth/*` (cadastro/login comprador).
- **Produtor:** `/api/producer/v1/auth/login` + onboarding; compliance gate em rotas protegidas.
- **Admin:** login em 2 etapas (senha → OTP e-mail) em `/api/admin/v1/auth/*`.

## 4. Módulos backend (domínio)

| Módulo | Responsabilidade |
| --- | --- |
| Events / Commercial | CRUD eventos, setores, lotes, readiness |
| Checkout / Payment | Reserva, Pix/cartão, tentativas |
| Tickets | Carteira, transferência, cancelamento |
| Operation | Check-in QR, facial, lista |
| Finance / Payouts | Ledger, repasse, freeze |
| Compliance | Termos versionados, `forceAcceptance` |
| KYC | Documentos produtor (titular) |
| Biometry | Enrollment e embeddings |

## 5. Integrações e dependências

| Sistema | Uso |
| --- | --- |
| MySQL | Persistência Prisma |
| Pagar.me | Captura e estornos |
| Brevo | E-mail transacional (OTP admin, convites) |
| pulse-face | Extração/validação facial |
| EAS / Vercel | Deploy apps e web |

## 6. Variáveis de ambiente (resumo)

Ver tabela completa em [api-endpoints.md](./api-endpoints.md#recursos) e `backend/README.md`.

| Categoria | Exemplos |
| --- | --- |
| Banco | `DATABASE_URL` |
| Auth | `BETTER_AUTH_SECRET`, `BETTER_AUTH_URL` |
| Pagamentos | `PAGARME_SECRET_KEY`, `PAYMENTS_ENABLED` (apps) |
| Facial | `PULSE_FACE_SERVICE_URL`, flags `FACIAL_*` |
| Admin | seed `PULSE_ADMIN` (`bun run seed:pulse-admin`) |

## 7. Referências cruzadas

- [README.md](./README.md) — índice funcional
- [api-endpoints.md](./api-endpoints.md) — catálogo HTTP
- [../product/technical/ARCHITECTURE_PRINCIPLES.md](../product/technical/ARCHITECTURE_PRINCIPLES.md)
""",
    )

    # API endpoints - from section 24
    api_blocks = section_blocks(blocks, "24. Catálogo", None)
    api_table = ""
    for kind, _, data in api_blocks:
        if kind == "t":
            api_table = table_md(data)  # type: ignore
            break
    write(
        OUT / "api-endpoints.md",
        f"""# Catálogo de endpoints API

> Escopo: contrato HTTP consolidado | Público: engenharia | Plataforma: backend Elysia | Última revisão: 2026-05-19

## Legenda de status

{legend}

## 1. Visão geral

- **OpenAPI interativo:** `GET /swagger` (somente rotas documentadas; B2C canônico em `/api/client/v1`).
- **Health:** `GET /health`
- Rotas legadas B2C na raiz espelham `/api/client/v1` — preferir o prefixo canônico em novos clientes.

## 2. Admin (`/api/admin/v1`)

Ver detalhes em [pulse-admin.md](./pulse-admin.md#3-módulos--funcionalidades).

## 3. Producer (`/api/producer/v1`)

Ver [app-produtor.md](./app-produtor.md) e [producer-web.md](./producer-web.md).

## 4. Client B2C (`/api/client/v1`)

Ver [app-client.md](./app-client.md) e [client-web.md](./client-web.md).

## 5. Tabela consolidada

{api_table}

## 6. Recursos

Infra mínima: processo Node (backend), MySQL, opcional `pulse-face`, deploys SSR/static para webs, builds EAS para apps.

| Categoria | Variáveis (sem valores secretos) |
| --- | --- |
| Banco | `DATABASE_URL` / `MYSQL_*` |
| Auth | `BETTER_AUTH_SECRET`, `BETTER_AUTH_URL` |
| HTTP | `PORT`, CORS dos frontends |
| Pagamentos | `PAGARME_SECRET_KEY`, `PAYMENTS_ENABLED` |
| E-mail | `BREVO_API_KEY`, `BREVO_SENDER_*`, `PRODUCER_WEB_URL` |
| Facial | `BIOMETRIC_*`, `FACIAL_*`, `PULSE_FACE_SERVICE_*` |
| KYC/Admin | `KYC_STORAGE_PATH`, seed admin |

## 7. Referências cruzadas

- Código: `backend/src/index.ts`, controllers em `backend/src/presentation/controllers/`
- [arquitetura.md](./arquitetura.md)
""",
    )

    # Pulse Admin
    admin_blocks = section_blocks(blocks, "21. Pulse Admin", "22. Arquitetura")
    hu_labels = {
        "21.1": ("3.1", "Autenticação e sessão (HU01)", "[IMPLEMENTADO]"),
        "21.2": ("3.2", "Produtoras e KYC (HU02)", "[IMPLEMENTADO]"),
        "21.3": ("3.3", "Visão e saúde do checkout (HU03)", "[IMPLEMENTADO]"),
        "21.4": ("3.4", "Financeiro — repasses e freeze (HU04)", "[IMPLEMENTADO]"),
        "21.5": ("3.5", "Central de estornos (HU05)", "[IMPLEMENTADO]"),
        "21.6": ("3.6", "Compliance e termos (HU06)", "[IMPLEMENTADO]"),
    }
    admin_body: list[str] = []
    current_section = ""
    for kind, level, data in admin_blocks:
        if kind == "t":
            admin_body.append("\n" + table_md(data) + "\n")
            continue
        if kind != "p" or level == 1:
            continue
        if level == 2:
            key = data.split()[0] if data else ""
            if key in hu_labels:
                num, title, st = hu_labels[key]
                admin_body.append(f"\n### {num} {title} — {st}\n")
                current_section = key
            continue
        if level == 0 and isinstance(data, str) and not data.startswith("Tabela consolidada"):
            admin_body.append(f"\n{data}\n")

    write(
        OUT / "pulse-admin.md",
        f"""# Pulse Admin (backoffice)

> Escopo: operação interna Pulse | Público: `PULSE_ADMIN` | Plataforma: Producer Web `/admin/*` + API `/api/admin/v1` | Última revisão: 2026-05-19

## Legenda de status

{legend}

## 1. Visão geral

Painel interno para operadores com papel **`PULSE_ADMIN`** (Operador Pulse). UI isolada em `producer-web` sob `/admin/*` (layout e sidebar próprios). Login unificado em `/login` com ramificação por role após OTP.

**Não confundir** com o portal da produtora ([producer-web.md](./producer-web.md)).

## 2. Autenticação e acesso

### 2.1 Auth 2FA (HU01) — [IMPLEMENTADO]

| Regra | Detalhe |
| --- | --- |
| MUST | Usuário sem `PULSE_ADMIN` recebe **403** em `/admin/*` e API admin |
| Login etapa 1 | `POST /api/admin/v1/auth/login` → `requiresOtp=true` (sem token) |
| Login etapa 2 | `POST /api/admin/v1/auth/login/verify-otp` → Bearer (~7 dias) |
| Sessão | `GET /auth/me`, `POST /auth/logout` |
| Middleware | `AdminAuthMiddleware` em todas as rotas exceto login/verify-otp |

## 3. Módulos / funcionalidades

{"".join(admin_body).strip()}

## 4. Permissões (RBAC nesta plataforma)

| Papel | Acesso Pulse Admin |
| --- | --- |
| `PULSE_ADMIN` | ✅ Total |
| `PRODUCER` / `PRODUCER_MANAGER` / `STAFF` / `PROMOTER` / `CLIENT` | ❌ |

Matriz completa: [RBAC.md](../RBAC.md).

## 5. Integrações e dependências

- API admin → MySQL (produtoras, payouts, refunds, compliance, KYC)
- Brevo para OTP de login
- Pagar.me para estornos (HU05)
- Compliance: publicação de termos bloqueia produtor/cliente via `TermsComplianceMiddleware` ([CHECKOUT_COMPLIANCE.md](../CHECKOUT_COMPLIANCE.md)); admin isento

## 6. Backlog / pendências

| Item | Status |
| --- | --- |
| Histórico métricas persistido (HU03b) | [PENDENTE] |
| Detalhe produtora / menu ações (HU02b) | [PARCIAL] |
| Export extrato admin (HU04b) | [PARCIAL] |
| Detalhe linha estorno (HU05b) | [PARCIAL] |
| Moderação/suspender evento global | [PENDENTE] |
| Antifraude/chargeback automático | [PENDENTE] |

## 7. Referências cruzadas

- [api-endpoints.md](./api-endpoints.md#2-admin-apadminv1)
- [producer-web.md](./producer-web.md) — mesmo deploy, rotas distintas
- [app-produtor.md](./app-produtor.md) — KYC titular espelha fila admin
""",
    )

    # Client ecosystem from section 14
    client_blocks = section_blocks(blocks, "14. Ecossistema", "15. RBAC")
    client_sections = []
    for kind, level, data in client_blocks:
        if kind == "p" and level == 2:
            client_sections.append(f"\n### {data.replace('14.', '').strip()}\n")
        elif kind == "p" and level == 0 and not str(data).startswith("14."):
            client_sections.append(data + "\n")

    b2c_blocks = section_blocks(blocks, "20. Módulos B2C", "21. Pulse Admin")
    b2c_table = ""
    for kind, _, data in b2c_blocks:
        if kind == "t":
            b2c_table = table_md(data)
            break

    write(
        OUT / "app-client.md",
        f"""# App Cliente (mobile B2C)

> Escopo: compra, carteira, facial, promoter | Público: `CLIENT`, `PROMOTER` | Plataforma: Expo `app-client/` | Última revisão: 2026-05-19

## Legenda de status

{legend}

## 1. Visão geral

Aplicativo mobile do comprador final: descoberta, checkout, pagamento, carteira de ingressos, cadastro facial e área do promoter. API canônica: **`/api/client/v1`**.

## 2. Autenticação e acesso

| Fluxo | Status | Regras |
| --- | --- | --- |
| Cadastro / login B2C | [IMPLEMENTADO] | Better Auth; papel base `CLIENT` |
| Compliance gate (HU06) | [IMPLEMENTADO] | Rotas autenticadas bloqueadas até aceitar termos `forceAcceptance` |
| Promoter | [IMPLEMENTADO] | `CLIENT` + membership `PROMOTER`; rotas `/promoter` |

Ver [CHECKOUT_COMPLIANCE.md](../CHECKOUT_COMPLIANCE.md).

## 3. Módulos / funcionalidades

### 3.1 Descoberta e vitrine — [IMPLEMENTADO]

- Feed, busca e detalhe de evento
- Seleção de lotes com regras de janela de venda
- MUST: vitrine não exige login de produtor

### 3.2 Checkout e pagamento — [PARCIAL]

| Regra | Detalhe |
| --- | --- |
| Reserva | **10 minutos** por pedido |
| Tentativas | Máx. **3** por pedido |
| Emissão | Ingresso só após status **`PAID`** |
| Flag demo | `PAYMENTS_ENABLED=false` → UI «Vendas em breve» (sem captura real) |

### 3.3 Carteira, facial e cancelamento — [IMPLEMENTADO]

- Meus ingressos / QR fallback
- Cadastro facial (`FACIAL_ENROLLMENT_V2`, `PULSE_FACE_EXTRACT`)
- Cancelamento: até **24h antes** do início; ticket não `USED` (`GetCancelEligibilityUseCase`)
- Alinhar copy de UI se ainda citar 48h

### 3.4 Área Promoter — [IMPLEMENTADO]

- Rotas `/promoter`: vendas e comissões
- MUST: Dono/Gestor convidam promoter com conta CLIENT existente ([RBAC.md](../RBAC.md))
- Promoter **não** acessa App Produtor nem Producer Web

### 3.5 VIP / Membership — [PENDENTE]

- VIP de lote (`isVip` no batch): [IMPLEMENTADO] no comercial
- Programa de assinatura recorrente: [PENDENTE] (tela VIP [PARCIAL] mock)

## 4. Permissões (RBAC nesta plataforma)

| Capacidade | CLIENT | Promoter |
| --- | --- | --- |
| Comprar / checkout | ✅ | ✅ |
| Carteira / facial | ✅ | ✅ |
| Comissões `/promoter` | ❌ | ✅ |
| Painel produtor | ❌ | ❌ |

## 5. Integrações e dependências

- Backend `/api/client/v1/*`, espelho legado raiz
- Pagar.me quando `PAYMENTS_ENABLED=true`
- pulse-face para biometria
- Regras globais: [global-business-rules.md](../produto/regras-negocio/global-business-rules.md)

## 6. Backlog / pendências

| Item | Status |
| --- | --- |
| Pagamentos reais em produção | [PARCIAL] — flag |
| Membership recorrente | [PENDENTE] |
| Paridade copy cancelamento 24h | [PARCIAL] |

## 7. Referências cruzadas

- [client-web.md](./client-web.md) — vitrine web; checkout web [PENDENTE]
- [app-produtor.md](./app-produtor.md) — emissão manual e Access
- Mapa B2C:

{b2c_table}
""",
    )

    write(
        OUT / "client-web.md",
        f"""# Client Web (site comprador)

> Escopo: vitrine pública e descoberta | Público: visitante / comprador | Plataforma: Next.js `client-web/` | Última revisão: 2026-05-19

## Legenda de status

{legend}

## 1. Visão geral

Site público de eventos: home, feed e detalhe. Seleção de lotes na web. **Checkout integrado na web:** [PENDENTE] — CTA direciona para App Cliente («Vendas em breve» quando pagamentos desabilitados).

## 2. Autenticação e acesso

| Fluxo | Status |
| --- | --- |
| Navegação anônima (vitrine) | [IMPLEMENTADO] |
| Login comprador | [PENDENTE] / via app |
| Carteira web | [PENDENTE] |

## 3. Módulos / funcionalidades

### 3.1 Vitrine e detalhe — [IMPLEMENTADO]

- Listagem/feed de eventos publicados
- Detalhe público com painel de lotes ([client-web/src/components/events/event-ticket-panel.tsx](../../client-web/src/components/events/event-ticket-panel.tsx))
- Exibe lotes à venda, esgotados e «em breve» conforme janela

### 3.2 Seleção de ingressos — [IMPLEMENTADO]

- Escolha de lotes e quantidades (respeita limite **4/evento/CPF** no backend no checkout)
- Sem finalização de pagamento na web

### 3.3 Checkout — [PENDENTE]

- Integração com `/api/client/v1/checkout` e pagamentos
- Compliance gate HU06 quando autenticado

### 3.4 Carteira e pós-compra — [PENDENTE]

- Meus ingressos, facial e cancelamento permanecem no App Cliente por enquanto

## 4. Permissões (RBAC nesta plataforma)

Apenas experiência **CLIENT** anônima ou futura autenticada. Papéis de produtora não aplicam.

## 5. Integrações e dependências

- API pública de eventos (`/api/client/v1/events` ou rotas públicas equivalentes)
- Mesmas regras de negócio de preço/lote que App Cliente

## 6. Backlog / pendências

| Item | Status |
| --- | --- |
| Checkout web completo | [PENDENTE] |
| Login + carteira web | [PENDENTE] |
| SSO / deep link para app | [PENDENTE] |

## 7. Referências cruzadas

- [app-client.md](./app-client.md)
- [README.md](./README.md#regras-transversais-todas-as-plataformas)
""",
    )

    # Producer epics for app-produtor
    epic_blocks = section_blocks(blocks, "5. Épico 1", "14. Ecossistema")
    epic_metas = parse_hu_meta(epic_blocks)

    def build_producer_app_md() -> str:
        modules = []
        epic_names = [
            ("5. Épico 1 — Acesso & Onboarding", "3.1 Acesso & Onboarding", "[IMPLEMENTADO]"),
            ("6. Épico 2 — Perfil & Configurações", "3.2 Perfil & Configurações", "[PARCIAL]"),
            ("7. Épico 3 — Gestão de Eventos", "3.3 Gestão de Eventos", "[IMPLEMENTADO]"),
            ("8. Épico 4 — Oferta Comercial", "3.4 Oferta Comercial", "[IMPLEMENTADO]"),
            ("9. Épico 5 — Dashboard do Produtor", "3.5 Dashboard", "[IMPLEMENTADO]"),
            ("10. Épico 6 — Pedidos & Participantes", "3.6 Pedidos & Participantes", "[IMPLEMENTADO]"),
            ("11. Épico 7 — Operação / Access", "3.7 Access (Check-in)", "[IMPLEMENTADO]"),
            ("12. Épico 8 — Financeiro & Repasses", "3.8 Financeiro", "[IMPLEMENTADO]"),
            ("13. Épico 9 — Roles", "3.9 Equipe & RBAC", "[PARCIAL]"),
        ]
        for epic_title, mod_title, status in epic_names:
            eb = section_blocks(blocks, epic_title.split("—")[0].strip(), None)
            # fix: use full start
            eb = section_blocks(
                blocks,
                epic_title[: epic_title.index("—")].strip() if "—" in epic_title else epic_title,
                None,
            )
            # simpler: slice by epic number
            num = epic_title.split(".")[0].strip()
            next_nums = [str(i) for i in range(int(num) + 1, 14)]
            eb = []
            cap = False
            for kind, level, data in blocks:
                if kind == "p" and level == 1 and isinstance(data, str):
                    if data.startswith(f"{num}. Épico"):
                        cap = True
                    elif cap and data.startswith(("14.", "15.")) or (
                        cap
                        and level == 1
                        and re.match(r"^\d+\.", data)
                        and not data.startswith(f"{num}.")
                    ):
                        if not data.startswith(f"{num}."):
                            break
                if cap:
                    eb.append((kind, level, data))

            epic_status = ""
            feat_table = ""
            hu_parts = []
            for kind, level, data in eb:
                if kind == "t" and not feat_table:
                    feat_table = table_md(data)  # type: ignore
                if kind == "p" and level == 2 and str(data).startswith("HU "):
                    sec = hu_from_backup(data)
                    meta = epic_metas.get(data, {})
                    hu_parts.append(render_hu(data, meta, sec))

            for kind, level, data in eb:
                if kind == "p" and level == 0 and "Status do épico" in str(data):
                    epic_status = data

            modules.append(
                f"### {mod_title} — {status}\n\n{epic_status}\n\n"
                + (f"{feat_table}\n\n" if feat_table else "")
                + "\n".join(hu_parts)
            )

        return "\n\n".join(modules)

    app_prod_body = build_producer_app_md()

    write(
        OUT / "app-produtor.md",
        f"""# App Produtor (mobile)

> Escopo: operação completa da produtora | Público: Dono (`PRODUCER`), Gestor (`PRODUCER_MANAGER`), Staff (`STAFF`) | Plataforma: Expo `app-producer/` | Última revisão: 2026-05-19

## Legenda de status

{legend}

## 1. Visão geral

Aplicativo móvel para o produtor operar o ciclo do evento: conta, eventos, lotes, participantes, **Access** (check-in), financeiro e equipe. API: **`/api/producer/v1`**.

**Promoter** usa preferencialmente o [App Cliente](./app-client.md). **Pulse Admin** é separado ([pulse-admin.md](./pulse-admin.md)).

### Decisões transversais (resumo)

| Tema | Regra |
| --- | --- |
| Publicação | Readiness mínima (oferta + dados) antes de `PUBLISHED` |
| Check-in | Facial principal; QR e lista como fallback; offline [PARCIAL] |
| Financeiro | Somente **Dono** na aba Finance global; Gestor vê por evento |
| Emissão manual | Fluxo auditável separado de venda plataforma |

## 2. Autenticação e acesso

### 2.1 Fluxo de login — [IMPLEMENTADO]

| Etapa | Comportamento |
| --- | --- |
| Login | E-mail + senha → `/api/producer/v1/auth/login` |
| `mustChangePassword` | Redireciona troca obrigatória |
| Termos pendentes | Bloqueia até aceite (`/compliance`) |
| Papéis | Dono, Gestor, Staff entram; `CLIENT` puro não entra |

**Critérios de aceite:** logout; recovery; termos bloqueiam; staff sem onboarding completo de conta titular.

## 3. Módulos / funcionalidades

{app_prod_body}

## 4. Permissões (RBAC nesta plataforma)

Resumo (detalhe em [RBAC.md](../RBAC.md)):

| Capacidade | Dono | Gestor | Staff |
| --- | --- | --- | --- |
| Tab Início (KPI empresa) | ✅ | ❌ | ❌ |
| Criar/publicar evento | ✅ | ✅ | ❌ |
| Tab Access | ✅ | ✅ | ✅ |
| Tab Finance global | ✅ | ❌ | ❌ |
| Financeiro por evento | ✅ | ✅ | ❌ |
| Convidar Gestor/Staff | ✅ | ❌ | ❌ |
| Convidar Promoter | ✅ | ✅ | ❌ |

## 5. Integrações e dependências

- Backend producer + operation (check-in)
- Câmera / facial / QR offline queue [PARCIAL]
- Repasse: ver [payout-policies.md](../produto/regras-negocio/payout-policies.md) — job `ReleaseRetainedPayoutsUseCase` (D+1 após término; legado 10 check-ins não implementado)

## 6. Backlog / pendências

| Item | Status |
| --- | --- |
| Insights preditivos dashboard | [PENDENTE] |
| Offline check-in robusto | [PARCIAL] |
| RBAC blindagem UI | [PARCIAL] |
| Despublicar/cancelar evento UX | [PARCIAL] |
| Encadear próximo lote | [CONFIRMAR COM PRODUTO] |

## 7. Referências cruzadas

- [producer-web.md](./producer-web.md) — paridade web
- [app-client.md](./app-client.md) — promoter
- [api-endpoints.md](./api-endpoints.md#3-producer-apiproducerv1)
""",
    )

    # Producer Web
    pw_blocks = section_blocks(blocks, "16. Producer Web", "17. Legado")
    pw_text = "\n".join(
        t for k, _, t in pw_blocks if k == "p" and _ == 0 and not t.startswith("16.")
    )

    write(
        OUT / "producer-web.md",
        f"""# Producer Web (portal da produtora)

> Escopo: painel web da produtora (não é Admin) | Público: Dono, Gestor, Staff | Plataforma: Next.js `producer-web/` rotas `(producer)/*` | Última revisão: 2026-05-19

## Legenda de status

{legend}

## 1. Visão geral

Portal **desktop-first** da produtora: dashboard, eventos, financeiro, equipe, onboarding e listas. Compartilha deploy com Pulse Admin, mas rotas **`/admin/*` estão excluídas** deste documento ([pulse-admin.md](./pulse-admin.md)).

{pw_text}

### Rotas principais (implementadas)

| Rota | Função | Status |
| --- | --- | --- |
| `/login`, `/forgot-password`, `/set-password` | Auth | [IMPLEMENTADO] |
| `/onboarding/*` | Cadastro produtora + KYC | [IMPLEMENTADO] |
| `/dashboard` | KPIs e atalhos | [IMPLEMENTADO] |
| `/events`, `/events/new`, `/events/[id]` | CRUD eventos | [IMPLEMENTADO] |
| `/finance`, `/finance/payouts`, `/finance/statement` | Financeiro | [PARCIAL] |
| `/team` | Equipe (gestor, staff, promoter) | [IMPLEMENTADO] |
| `/settings` | Perfil e conta | [PARCIAL] |
| `/lists` | Listas operacionais | [PARCIAL] |
| `/vip` | Programa VIP | [PARCIAL] stub «em breve» |

## 2. Autenticação e acesso

- Login produtor: `POST /api/producer/v1/auth/login`
- Compliance gate igual app ([CHECKOUT_COMPLIANCE.md](../CHECKOUT_COMPLIANCE.md))
- Middleware web restringe sidebar por papel (Gestor sem `/finance` global)

## 3. Módulos / funcionalidades

### 3.1 Onboarding & KYC — [IMPLEMENTADO]

Fluxo `/onboarding/*` + upload documentos titular; fila revisada no [Pulse Admin](./pulse-admin.md).

### 3.2 Eventos & oferta comercial — [IMPLEMENTADO]

- CRUD e publicação com readiness
- Comercial/lotes: paridade com app; gestão avançada de lotes [PARCIAL] vs App Produtor
- Criação rápida sem setores na web em alguns fluxos [PARCIAL]

### 3.3 Dashboard — [IMPLEMENTADO]

Cards, gráfico de vendas, eventos próximos, alertas. Insights preditivos [PENDENTE].

### 3.4 Participantes & emissão manual — [IMPLEMENTADO]

Lista, busca, cortesia e venda direta. Exportações [PENDENTE].

### 3.5 Check-in ao vivo — [PENDENTE]

Paridade com App Produtor Access: botões «Em breve» no dashboard; API `operation/*` existe no backend.

### 3.6 Financeiro — [PARCIAL]

| Submódulo | App Produtor | Producer Web |
| --- | --- | --- |
| Resumo / por evento | [IMPLEMENTADO] | [IMPLEMENTADO] / [PARCIAL] |
| Cancelamentos na UI | [IMPLEMENTADO] | [PENDENTE] |
| Repasses / antecipação | [IMPLEMENTADO] | [IMPLEMENTADO] |
| Comissões promoter | — | [PENDENTE] |

### 3.7 Equipe & RBAC — [PARCIAL]

Convites: `invite-manager`, staff, promoter ([RBAC.md](../RBAC.md)). Validação visual por tela em evolução.

## 4. Permissões (RBAC nesta plataforma)

Área produtora (`/dashboard`, `/events`, …) — sem acesso `/admin/*`.

| Capacidade | Dono | Gestor | Staff |
| --- | --- | --- | --- |
| Dashboard GMV empresa | ✅ | ❌ | ❌ |
| `/finance` global | ✅ | ❌ | ❌ |
| Financeiro por evento | ✅ | ✅ | ❌ |
| Equipe — convidar Gestor/Staff | ✅ | ❌ | ❌ |
| Equipe — convidar Promoter | ✅ | ✅ | ❌ |

## 5. Integrações e dependências

- Eden Treaty → backend `/api/producer/v1`
- Better Auth cliente, TanStack Query
- Roadmap histórico: [ROADMAP-PRODUCER-WEB.md](../ROADMAP-PRODUCER-WEB.md)

## 6. Backlog / pendências

| Item | Status |
| --- | --- |
| Check-in web ao vivo | [PENDENTE] |
| Checkout/carteira (escopo B2C no client-web) | N/A |
| Telas cancelamentos/comissões financeiro | [PENDENTE] |
| Analytics preditivos | [PENDENTE] |

## 7. Referências cruzadas

- [app-produtor.md](./app-produtor.md)
- [pulse-admin.md](./pulse-admin.md)
- [client-web.md](./client-web.md)
""",
    )

    print("Generated:", list(OUT.glob("*.md")))


if __name__ == "__main__":
    main()
