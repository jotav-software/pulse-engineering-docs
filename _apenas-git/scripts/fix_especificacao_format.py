#!/usr/bin/env python3
"""Correções de formatação em docs/especificacao_funcional.docx — sem alterar texto substantivo."""

from __future__ import annotations

import hashlib
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
DOC_PATH = ROOT / "docs/especificacao_funcional.docx"


def paragraph_texts(doc: Document) -> list[str]:
    return [(p.text or "") for p in doc.paragraphs]


def texts_fingerprint(texts: list[str]) -> str:
    joined = "\n".join(texts)
    return hashlib.sha256(joined.encode("utf-8")).hexdigest()


def fix_table_grids(doc: Document) -> int:
    fixed = 0
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        ncells = len(tbl.rows[0].cells)
        grid = tbl._tbl.find(qn("w:tblGrid"))
        if grid is None:
            continue
        cols = grid.findall(qn("w:gridCol"))
        if len(cols) == ncells:
            continue
        default_w = 1800
        if cols:
            w = cols[0].get(qn("w:w"))
            if w:
                default_w = max(800, int(int(w) * len(cols) / max(ncells, 1)))
        for _ in range(ncells - len(cols)):
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(default_w))
            grid.append(gc)
        fixed += 1
    return fixed


def fix_strikethrough_runs(doc: Document) -> int:
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            r_pr = run._element.rPr
            if r_pr is None:
                continue
            strike = r_pr.find(qn("w:strike"))
            if strike is None:
                continue
            val = strike.get(qn("w:val"))
            if val in ("0", "false", "off"):
                run.font.strike = True
                count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        r_pr = run._element.rPr
                        if r_pr is None:
                            continue
                        strike = r_pr.find(qn("w:strike"))
                        if strike is None:
                            continue
                        val = strike.get(qn("w:val"))
                        if val in ("0", "false", "off"):
                            run.font.strike = True
                            count += 1
    return count


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_duplicate_legend(doc: Document) -> bool:
    """Remove segunda ocorrência idêntica da legenda (parágrafo após H1 seção 1)."""
    if len(doc.paragraphs) < 10:
        return False
    p1 = (doc.paragraphs[1].text or "").strip()
    p8 = (doc.paragraphs[8].text or "").strip()
    if p1 and p1 == p8:
        delete_paragraph(doc.paragraphs[8])
        return True
    return False


def collapse_consecutive_empty_paragraphs(doc: Document) -> int:
    removed = 0
    i = 0
    while i < len(doc.paragraphs):
        if not (doc.paragraphs[i].text or "").strip():
            j = i + 1
            while j < len(doc.paragraphs) and not (doc.paragraphs[j].text or "").strip():
                delete_paragraph(doc.paragraphs[j])
                removed += 1
            i += 1
        else:
            i += 1
    return removed


def fix_section_21_order(doc: Document) -> bool:
    """21.7 antes de 21.8; texto introdutório imediatamente antes da tabela HU."""
    idx_218 = idx_217 = idx_intro = tbl_hu = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t == "21.8 Tabela HU × rota × status":
            idx_218 = i
        elif t == "21.7 Mapa HU × rota × status":
            idx_217 = i
        elif t.startswith("Tabela consolidada na seção 21.8"):
            idx_intro = i
    for ti, tbl in enumerate(doc.tables):
        if not tbl.rows:
            continue
        hdr = [c.text.strip() for c in tbl.rows[0].cells]
        if hdr[:4] == ["HU", "Fluxo", "Rota/UI", "Status"]:
            tbl_hu = tbl
            break
    if idx_218 is None or idx_217 is None or idx_intro is None or tbl_hu is None:
        return False

    body = doc.element.body
    el_218 = doc.paragraphs[idx_218]._element
    el_217 = doc.paragraphs[idx_217]._element
    el_intro = doc.paragraphs[idx_intro]._element
    el_tbl = tbl_hu._tbl

    # Garantir ordem: 21.7, 21.8, intro, tabela
    if el_217.getparent() is body:
        body.remove(el_217)
        el_218.addprevious(el_217)

    if el_tbl.getparent() is body:
        body.remove(el_tbl)
        el_intro.addnext(el_tbl)

    return True


def main() -> int:
    path = DOC_PATH
    if not path.exists():
        print(f"Arquivo não encontrado: {path}", file=sys.stderr)
        return 1

    doc = Document(str(path))
    before_texts = paragraph_texts(doc)
    before_fp = texts_fingerprint(before_texts)

    grids_fixed = fix_table_grids(doc)
    strikes_fixed = fix_strikethrough_runs(doc)
    legend_removed = remove_duplicate_legend(doc)
    empties_removed = collapse_consecutive_empty_paragraphs(doc)
    sec21_fixed = fix_section_21_order(doc)

    after_texts = paragraph_texts(doc)
    after_fp = texts_fingerprint(after_texts)

  # Texto substantivo: mesmas linhas não vazias na mesma ordem (exceto legenda duplicada removida)
    before_nonempty = [t for t in before_texts if t.strip()]
    after_nonempty = [t for t in after_texts if t.strip()]
    legend_dup = (
        len(before_nonempty) - len(after_nonempty) == 1
        and before_nonempty.count(before_nonempty[1]) >= 2
    )
    if before_nonempty != after_nonempty and not legend_dup:
        # Permitir apenas remoção da legenda duplicada
        if not (
            len(before_nonempty) == len(after_nonempty) + 1
            and after_nonempty == [t for t in before_nonempty if t != before_nonempty[1] or before_nonempty.index(t) != 1]
        ):
            b_set = before_nonempty
            a_set = after_nonempty
            if b_set != a_set:
                print("ERRO: texto substantivo alterado!", file=sys.stderr)
                print(f"  antes: {len(before_nonempty)} linhas, depois: {len(after_nonempty)}")
                for i, (a, b) in enumerate(zip(before_nonempty[:5], after_nonempty[:5])):
                    if a != b:
                        print(f"  diff@{i}: {a[:60]!r} vs {b[:60]!r}")
                return 2

    doc.save(str(path))

    print("Correções aplicadas:")
    print(f"  Tabelas com tblGrid corrigido: {grids_fixed}")
    print(f"  Runs com strikethrough corrigido: {strikes_fixed}")
    print(f"  Legenda duplicada removida: {legend_removed}")
    print(f"  Parágrafos vazios consecutivos removidos: {empties_removed}")
    print(f"  Seção 21 reordenada: {sec21_fixed}")
    print(f"  Fingerprint antes: {before_fp[:16]}...")
    print(f"  Fingerprint depois: {texts_fingerprint(paragraph_texts(doc))[:16]}...")
    print(f"  Parágrafos: {len(before_texts)} -> {len(after_texts)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
