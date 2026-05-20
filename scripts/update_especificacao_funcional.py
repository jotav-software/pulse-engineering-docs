#!/usr/bin/env python3
"""Atualiza docs/especificacao_funcional.docx — fonte única de verdade do ecossistema Pulse."""

from __future__ import annotations

import re
import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
SRC_PRODUCER = ROOT / "docs/especificacao_funcional_mvp_app_produtor_pulse_pro.docx"
OUT = ROOT / "docs/especificacao_funcional.docx"
TODAY = date.today()

LEGENDA = (
    f"Legenda de status (de/para código + docs, atualizado em {TODAY.isoformat()}): "
    "[IMPLEMENTADO] entregue e utilizável | "
    "[PARCIAL] fluxo existe com lacunas conhecidas (inclui UI marcada «em breve») | "
    "[PENDENTE] não entregue ou só especificado | "
    "[CONFIRMAR COM PRODUTO] evidência insuficiente — validar com produto."
)

GLOSSARIO = (
    "Glossário de plataformas: Produtor App = App Produtor (mobile) | "
    "Producer Web = portal web da produtora + Pulse Admin (`/admin/*`) | "
    "App Cliente = App Client (mobile B2C) | Client Web = site comprador | "
    "Pulse Admin = papel `PULSE_ADMIN` (Operador Pulse), não é dono de produtora."
)


def set_strike(run, strike: bool = True) -> None:
    run.font.strike = strike


def strike_paragraph(para) -> None:
    for run in para.runs:
        set_strike(run, True)


def replace_in_paragraph(para, replacements: dict[str, str]) -> None:
    full = para.text
    if not full.strip():
        return
    new = full
    for old, new_val in replacements.items():
        new = new.replace(old, new_val)
    if new != full:
        para.clear()
        para.add_run(new)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_paragraph(doc: Document, prefix: str) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(prefix):
            return i
    return None


def remove_section_range(doc: Document, start_prefix: str, end_prefix: str) -> None:
    """Remove parágrafos de start_prefix (inclusive) até antes de end_prefix."""
    start = find_paragraph(doc, start_prefix)
    end = find_paragraph(doc, end_prefix)
    if start is None:
        return
    end_idx = end if end is not None else len(doc.paragraphs)
    to_remove = list(doc.paragraphs[start:end_idx])
    for p in to_remove:
        delete_paragraph(p)


def insert_blocks_before(anchor, blocks: list[tuple[str, str]]) -> None:
    """Insere blocos na ordem de leitura (primeiro bloco no topo)."""
    for style_name, text in blocks:
        new_p = anchor.insert_paragraph_before(text)
        try:
            new_p.style = style_name
        except Exception:
            new_p.style = "normal"


def add_table_from_rows(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def insert_table_after_paragraph(
    doc: Document, paragraph, headers: list[str], rows: list[list[str]]
) -> None:
    """Cria tabela e posiciona logo após o parágrafo âncora."""
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._element.addnext(tbl)


# --- Substituições globais ---
STATUS_REPLACEMENTS = {
    "Status do épico: ⏳ Não implementado / especificado — Dar visão": (
        "Status do épico: [IMPLEMENTADO] Produtor App e Producer Web | "
        "[PARCIAL] insights avançados — Dar visão"
    ),
    "Status: ⏳ Não implementado\n": "Status: [IMPLEMENTADO] Produtor App e Producer Web\n",
    "Status do épico: ⏳ Não implementado — Dar controle operacional": (
        "Status do épico: [IMPLEMENTADO] Produtor App e Producer Web (participantes/emissão manual) | "
        "[PARCIAL] exportações — Dar controle operacional"
    ),
    "Status do épico: ⏳ Não implementado — Executar a entrada": (
        "Status do épico: [IMPLEMENTADO] Produtor App | [PENDENTE] Producer Web (check-in ao vivo) — Executar a entrada"
    ),
    "HU 7.1 — Access operacional do evento\nStatus: ⏳ Não implementado": (
        "HU 7.1 — Access operacional do evento\nStatus: [IMPLEMENTADO] Produtor App | [PENDENTE] Producer Web"
    ),
    "Status do épico: ⏳ Não implementado / especificado com cuidado — Dar ao produtor": (
        "Status do épico: [IMPLEMENTADO] Produtor App | [PARCIAL] Producer Web (sem telas de cancelamentos/comissões) — Dar ao produtor"
    ),
    "HU 8.1 — Painel financeiro do produtor e repasses\nStatus: ⏳ Não implementado": (
        "HU 8.1 — Painel financeiro do produtor e repasses\nStatus: [IMPLEMENTADO] Produtor App | [PARCIAL] Producer Web"
    ),
    "Status do épico: ✅ Base implementada / ⏳ avançados pendentes — Estruturar": (
        "Status do épico: [IMPLEMENTADO] base e gestão avançada de lotes no Produtor App | "
        "[PARCIAL] Producer Web — Estruturar"
    ),
    "HU 4.1 — Configurar oferta comercial do evento\nStatus: ✅ Base implementada / ⏳ avançados pendentes": (
        "HU 4.1 — Configurar oferta comercial do evento\nStatus: [IMPLEMENTADO] reordenar/duplicar/pausar no Produtor App | [PARCIAL] Producer Web"
    ),
    "Admin da plataforma fica fora do escopo detalhado deste documento e entra apenas como papel futuro.": (
        "Pulse Admin (Producer Web `/admin/*`, papel `PULSE_ADMIN`) está documentado na seção 21; "
        "não substitui regras B2C nem portal da produtora."
    ),
    "Legenda de status: ✅ Implementado   🟡 Parcial / em refinamento   ⏳ Não implementado   🔵 Futuro": LEGENDA,
    "Este documento resume o app do produtor": (
        "Este documento consolida regras do ecossistema Pulse (Produtor App, Producer Web, "
        "App Cliente, Client Web e backend)"
    ),
}

STRIKE_CONTAINS = [
    "Gestão avançada de lotes | ⏳ | Especificado; não implementado",
    "Dashboard do Produtor | ⏳ | Especificado; implementação pendente",
    "Pedidos & Participantes | ⏳ | Especificado; implementação pendente",
    "Operação / Access | ⏳ | Especificado; implementação pendente",
    "Financeiro & Repasses | ⏳ | Especificação consolidada; requer validação técnica",
    "Promoter no app cliente | ⏳ | Área Minhas Vendas / Promoter pendente",
    "Reordenar (pendente)",
    "Duplicar (pendente)",
    "Pausar/Reativar (pendente)",
]

HU91_STRIKE = (
    "Staff pode fazer tudo que o promoter faz no contexto comercial, "
    "desde que o produtor permita."
)
HU91_REPLACE = (
    "Staff limita-se à operação de porta (QR, facial, lista). Não acessa comercial nem comissões. "
    "Promoter atua apenas no App Cliente (`/promoter`). Ver docs/RBAC.md."
)

HU91_ADMIN_STRIKE = "Admin da plataforma fica como papel futuro nesta HU."
HU91_ADMIN_REPLACE = (
    "Pulse Admin (`PULSE_ADMIN`) está [IMPLEMENTADO] — ver seção 21. "
    "Moderação global de eventos e antifraude automático permanecem [PENDENTE]."
)


def patch_roadmap_table(table) -> None:
    rows_data = [
        ("1", "Acesso & Onboarding", "Entrada segura e configuração inicial", "[IMPLEMENTADO]", "App + Web produtor"),
        ("2", "Perfil & Configurações", "Conta, banco, defaults, equipe", "[PARCIAL]", "Banco: load parcial no app; termos/exclusão mais completos na web"),
        ("3", "Gestão de Eventos", "CRUD e publicação", "[IMPLEMENTADO]", "Sem delete de evento em nenhuma plataforma"),
        ("4", "Oferta Comercial", "Setores/lotes e readiness", "[IMPLEMENTADO]", "App completo; web sem setores na criação rápida"),
        ("5", "Gestão Avançada de Lotes", "Reordenar, duplicar, pausa, encadeamento", "[IMPLEMENTADO]", "Produtor App; [PARCIAL] Producer Web"),
        ("6", "Dashboard do Produtor", "KPIs, alertas, gráfico", "[IMPLEMENTADO]", "Insights avançados [PENDENTE]"),
        ("7", "Pedidos & Participantes", "Lista, emissão manual, cortesia", "[IMPLEMENTADO]", "App + Web; export [PENDENTE]"),
        ("8", "Operação / Access", "QR, facial, lista, offline", "[IMPLEMENTADO]", "App | Web check-in ao vivo [PENDENTE]"),
        ("9", "Financeiro & Repasses", "Ledger, repasse, antecipação", "[IMPLEMENTADO]", "App completo; web sem cancelamentos UI"),
        ("10", "Roles, Equipe & Promoter", "RBAC + promoter", "[PARCIAL]", "Alvo em docs/RBAC.md; matriz em evolução"),
        ("11", "Pulse Admin + evoluções", "Painel interno + analytics", "[PARCIAL]", "Admin core [IMPLEMENTADO]; analytics/moderação [PENDENTE]"),
    ]
    for i, row_vals in enumerate(rows_data, start=1):
        if i >= len(table.rows):
            break
        row = table.rows[i]
        for j, val in enumerate(row_vals):
            if j < len(row.cells):
                row.cells[j].text = val


def _row_add_cell(row) -> None:
    tc = OxmlElement("w:tc")
    row._tr.append(tc)


def add_platform_columns_to_feature_table(
    table, platform_status: list[tuple[str, str, str, str, str, str]]
) -> None:
    if len(table.rows) == 0:
        return
    header = table.rows[0]
    if any("Produtor App" in c.text for c in header.cells):
        return
    for ri in range(len(table.rows)):
        row = table.rows[ri]
        while len(row.cells) < 6:
            _row_add_cell(row)
        if ri == 0:
            labels = [
                "Feature",
                "Status geral",
                "Produtor App",
                "Producer Web",
                "Client App",
                "Client Web",
            ]
            for ci, lab in enumerate(labels):
                row.cells[ci].text = lab
        elif ri - 1 < len(platform_status):
            vals = list(platform_status[ri - 1])
            for ci, v in enumerate(vals):
                row.cells[ci].text = v


FEATURE_TABLE_PATCHES: dict[int, list[tuple[str, str, str, str, str, str]]] = {
    6: [
        ("Reordenar", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[PARCIAL]", "—", "—"),
        ("Duplicar", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[PARCIAL]", "—", "—"),
        ("Pausar/Reativar", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[PARCIAL]", "—", "—"),
        ("Editar lote com venda ativa", "[PARCIAL]", "[PARCIAL]", "[PARCIAL]", "—", "—"),
        ("Encadear próximo lote", "[PARCIAL]", "[PARCIAL]", "[CONFIRMAR COM PRODUTO]", "—", "—"),
    ],
    7: [
        ("Cards globais", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "—", "—"),
        ("Gráfico de vendas", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "—", "—"),
        ("Eventos próximos", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "—", "—"),
        ("Alertas operacionais", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "—", "—"),
        ("Insights preditivos", "[PENDENTE]", "[PENDENTE]", "[PENDENTE]", "—", "—"),
    ],
    8: [
        ("Listagem de participantes", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "—", "—"),
        ("Busca manual", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "—", "—"),
        ("Emissão manual", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "—", "—"),
        ("Cortesia", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "—", "—"),
        ("Venda direta do produtor", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "—", "—"),
    ],
    9: [
        ("Tela Access", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[PENDENTE]", "—", "—"),
        ("Scanner QR", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[PENDENTE]", "—", "—"),
        ("Facial", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[PENDENTE]", "—", "—"),
        ("Lista e busca manual", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[PARCIAL]", "—", "—"),
        ("Status do ingresso", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[PARCIAL]", "—", "—"),
        ("Contingência offline", "[PARCIAL]", "[PARCIAL]", "[PENDENTE]", "—", "—"),
    ],
    10: [
        ("Resumo financeiro", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "—", "—"),
        ("Detalhe por evento", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[PARCIAL]", "—", "—"),
        ("Cancelamentos/estornos", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[PENDENTE]", "—", "—"),
        ("Repasse e bloqueios", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "—", "—"),
        ("Conciliação e auditoria", "[PARCIAL]", "[PARCIAL]", "[PARCIAL]", "—", "—"),
    ],
    11: [
        ("Matriz de permissões", "[PARCIAL]", "[PARCIAL]", "[PARCIAL]", "—", "—"),
        ("Gestão de equipe por escopo", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "—", "—"),
        ("Experiência comercial do Staff", "[IMPLEMENTADO]", "[PARCIAL]", "—", "—", "—"),
        ("Experiência do Promoter no App Cliente", "[IMPLEMENTADO]", "—", "—", "[IMPLEMENTADO]", "—"),
    ],
}


def ecosystem_blocks() -> list[tuple[str, str]]:
    return [
        ("Heading 1", "14. Ecossistema — Cliente (App Cliente + Client Web)"),
        (
            "normal",
            "Escopo B2C (comprador e promoter). Referências: "
            "docs/product/especificacao_funcional_mvp_ingressos.docx, "
            "product/global-business-rules.md, product/checkout-compliance.md (pulse-engineering-docs).",
        ),
        ("Heading 2", "14.1 Descoberta e vitrine"),
        (
            "normal",
            "Client Web: home, feed e detalhe público do evento [IMPLEMENTADO]. "
            "Seleção de lotes na web [IMPLEMENTADO]. Checkout integrado na web [PENDENTE] — "
            "CTA direciona para App Cliente. "
            "App Cliente: feed, busca, detalhe [IMPLEMENTADO]. "
            "MUST: vitrine pública não exige login de produtor.",
        ),
        ("Heading 2", "14.2 Checkout, pagamento e carteira (App Cliente)"),
        (
            "normal",
            "MUST: reserva de estoque por 10 minutos; máximo 3 tentativas de pagamento por pedido; "
            "ingresso só emitido após status PAID (product/global-business-rules.md). "
            "App Cliente: fluxo implementado [PARCIAL] — flag PAYMENTS_ENABLED=false em demo impede pagamento real. "
            "Client Web: sem checkout [PENDENTE]. "
            "Gate HU06: rotas autenticadas bloqueadas até aceitar termos com forceAcceptance "
            "(PULSE_ADMIN isento — docs/CHECKOUT_COMPLIANCE.md).",
        ),
        ("Heading 2", "14.3 Ingressos, facial e cancelamento"),
        (
            "normal",
            "Carteira / meus ingressos no App Cliente [IMPLEMENTADO]. "
            "Cadastro facial [IMPLEMENTADO] com flags (FACIAL_ENROLLMENT_V2, PULSE_FACE_EXTRACT). "
            "MUST: cancelamento pelo comprador até 24h antes do início do evento, ticket não utilizado "
            "(GetCancelEligibilityUseCase — alinhar copy do app que ainda cita 48h). "
            "QR como fallback operacional [IMPLEMENTADO]. Client Web carteira [PENDENTE].",
        ),
        ("Heading 2", "14.4 Área Promoter (App Cliente)"),
        (
            "normal",
            "Rotas /promoter (vendas, comissões) [IMPLEMENTADO] para membership PROMOTER. "
            "MUST: Dono e Gestor convidam promoter com conta CLIENT já existente (docs/RBAC.md). "
            "Promoter não acessa Produtor App nem Producer Web.",
        ),
        ("Heading 2", "14.5 VIP / Membership"),
        (
            "normal",
            "VIP de lote (`isVip` no batch) [IMPLEMENTADO] no comercial. "
            "Programa de assinatura / membership recorrente [PENDENTE]: "
            "App Cliente tela VIP [PARCIAL] (mock); Producer Web /vip [PARCIAL] («em breve»); "
            "sem plano recorrente no schema.",
        ),
        ("Heading 1", "15. RBAC unificado (referência docs/RBAC.md)"),
        (
            "normal",
            "Papéis: CLIENT, PRODUCER (Dono), PRODUCER_MANAGER (Gestor), STAFF, PROMOTER, PULSE_ADMIN. "
            "Fonte canônica de permissões: docs/RBAC.md (2026-05-19). "
            "MUST: permissão validada no backend, nunca só na UI.",
        ),
        ("Heading 2", "15.1 Matriz resumida por aplicação"),
        (
            "normal",
            "Produtor App: épicos 5–13. Producer Web produtora: paridade parcial; check-in ao vivo [PENDENTE] "
            "(atalhos «Em breve» no dashboard). Client Web: vitrine [IMPLEMENTADO]; auth/checkout B2C [PENDENTE] "
            "(hoje usa API produtor — migrar para /api/client/v1). Client App: compra, carteira, promoter [IMPLEMENTADO]. "
            "Pulse Admin: seção 21.",
        ),
        ("Heading 1", "16. Producer Web — portal da produtora"),
        (
            "normal",
            "Rotas: /dashboard, /events, /finance/*, /team, /settings, onboarding /onboarding/*, /lists. "
            "Check-in operacional ao vivo [PENDENTE] (botão «Em breve» em quick-actions). "
            "Listas /lists: consulta participantes [PARCIAL]. "
            "Financeiro web: repasse e KPIs [IMPLEMENTADO]; cancelamentos/comissões UI [PENDENTE]. "
            "Área admin isolada em /admin/* — ver seção 21 (não confundir com portal produtor).",
        ),
        ("Heading 1", "17. Legado — numeração dos épicos produtor"),
        (
            "normal",
            "Seções 5–13 mantêm HUs do Produtor App. Status atualizados por de/para com código; "
            "consulte seções 14–16 para demais plataformas e seção 21 para Pulse Admin.",
        ),
    ]


def pulse_admin_blocks() -> list[tuple[str, str]]:
    return [
        ("Heading 1", "21. Pulse Admin — especificação operacional"),
        (
            "normal",
            "Painel interno Pulse para operadores com papel `PULSE_ADMIN` (Operador Pulse). "
            "UI: Producer Web em `/admin/*` (layout e sidebar isolados — HU01 isolamento de código). "
            "API: `/api/admin/v1/*`. Login unificado em `/login` com ramificação por role após OTP.",
        ),
        ("Heading 2", "21.1 Autenticação e sessão (HU01)"),
        (
            "normal",
            "MUST: usuário sem role PULSE_ADMIN não acessa `/admin/*` nem API admin (403). "
            "Given credenciais válidas When POST /api/admin/v1/auth/login Then requiresOtp=true (sem token). "
            "Given OTP válido When POST /api/admin/v1/auth/login/verify-otp Then token Bearer (~7 dias) e sessão admin. "
            "Rotas: logout, GET /auth/me. Middleware AdminAuthMiddleware em todas as rotas exceto login/verify-otp.",
        ),
        ("Heading 2", "21.2 Produtoras — onboarding e KYC (HU02)"),
        (
            "normal",
            "Tela /admin/produtoras: listagem com GMV 30d, busca, drawer criar produtora (CNPJ, taxa pulseFeeBps), "
            "reset de senha. API: GET/POST /producers, POST /producers/:id/reset-password. "
            "Subfluxo KYC titular: /admin/compliance/kyc — fila, aprovar, rejeitar, download documento. "
            "API KYC: GET /kyc/queue, GET /kyc/documents/:id, approve, reject, download.",
        ),
        ("Heading 2", "21.3 Visão geral e saúde do checkout (HU03)"),
        (
            "normal",
            "Tela /admin/visao: KPIs tráfego checkout 24h, latência p95, health gateways. "
            "API: GET /api/admin/v1/metrics/health. "
            "[PARCIAL] M1 — métricas em memória (MetricsStore), sem persistência histórica longa.",
        ),
        ("Heading 2", "21.4 Financeiro — repasses e freeze (HU04)"),
        (
            "normal",
            "Tela /admin/financeiro: abas repasses pendentes, congelados, liberados (30d); KPIs; "
            "modais freeze/unfreeze com motivo obrigatório (mín. 10 caracteres). "
            "MUST: evento congelado bloqueia saque do produtor (mesma regra do portal produtor). "
            "API: GET /payouts, GET /payouts/stats, POST /payouts/events/:eventId/freeze|unfreeze. "
            "UI secundária «Exportar extrato» / KPI chargeback [PARCIAL] — marcada «em breve».",
        ),
        ("Heading 2", "21.5 Central de estornos (HU05)"),
        (
            "normal",
            "Mesma tela financeiro: listagem estornos, busca pedido, validação e processamento via gateway. "
            "API: GET /refunds, /refunds/stats, /refunds/producers/:id/events, /refunds/search-orders, "
            "POST /refunds/validate, POST /refunds (executar). "
            "Detalhe linha e ações extras [PARCIAL] («Detalhes em breve»).",
        ),
        ("Heading 2", "21.6 Compliance e termos legais (HU06)"),
        (
            "normal",
            "Tela /admin/compliance: documentos versionados, publicar nova versão com forceAcceptance. "
            "API: GET /compliance, POST /compliance/documents. "
            "Produtor e cliente bloqueados por TermsComplianceMiddleware até aceitar; PULSE_ADMIN isento. "
            "Ver docs/CHECKOUT_COMPLIANCE.md.",
        ),
        ("Heading 2", "21.7 Mapa HU × rota × status"),
        (
            "normal",
            "Tabela consolidada na seção 21.8 (abaixo). UI «em breve» = [PARCIAL].",
        ),
    ]


def pulse_admin_hu_table_rows() -> list[list[str]]:
    return [
        ("HU01", "Auth 2FA + isolamento /admin", "/admin/*, /api/admin/v1/auth/*", "[IMPLEMENTADO]"),
        ("HU02", "Produtoras + KYC titular", "/admin/produtoras, /admin/compliance/kyc", "[IMPLEMENTADO]"),
        ("HU02b", "Detalhe produtora / ações menu", "producers-table", "[PARCIAL] em breve"),
        ("HU03", "Visão checkout 24h", "/admin/visao", "[IMPLEMENTADO]"),
        ("HU03b", "Histórico métricas persistido", "metrics", "[PENDENTE]"),
        ("HU04", "Repasses + freeze", "/admin/financeiro", "[IMPLEMENTADO]"),
        ("HU04b", "Export extrato admin", "financeiro-view", "[PARCIAL] em breve"),
        ("HU05", "Estornos centralizados", "/admin/financeiro", "[IMPLEMENTADO]"),
        ("HU05b", "Detalhe estorno linha", "refunds-table", "[PARCIAL] em breve"),
        ("HU06", "Compliance / termos", "/admin/compliance", "[IMPLEMENTADO]"),
        ("—", "Moderação/suspender evento global", "—", "[PENDENTE]"),
        ("—", "Antifraude/chargeback automático", "—", "[PENDENTE]"),
    ]


def post_mvp_blocks() -> list[tuple[str, str]]:
    return [
        ("Heading 1", "18. Evoluções pós-MVP e backlog estratégico"),
        (
            "normal",
            "Separar o que já existe do backlog. Pulse Admin **core** (seção 21) está [IMPLEMENTADO]; "
            "itens abaixo permanecem pós-MVP salvo indicação em contrário.",
        ),
        ("Heading 2", "18.1 Já entregue (não tratar como futuro)"),
        (
            "normal",
            "Pulse Admin: visão, produtoras, financeiro admin (freeze/estornos), compliance/KYC. "
            "Promoter no App Cliente. Gestão avançada de lotes no Produtor App. "
            "Dashboard e financeiro produtor no App.",
        ),
        ("Heading 2", "18.2 Backlog pós-MVP"),
        (
            "normal",
            "Analytics preditivos e precificação automática [PENDENTE]. "
            "Check-in ao vivo na Producer Web (paridade com Produtor App) [PENDENTE]. "
            "Checkout e carteira no Client Web com /api/client/v1 [PENDENTE]. "
            "Membership/VIP recorrente (assinatura) [PENDENTE]. "
            "Moderação global de eventos e antifraude admin [PENDENTE]. "
            "Operação Access offline com resolução de conflito avançada [PARCIAL]. "
            "Exportações em massa (participantes, financeiro) [PENDENTE].",
        ),
    ]


def appendix_architecture(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("22. Arquitetura do sistema (visão de alto nível)", level=1)
    doc.add_paragraph(
        "Ecossistema monorepo com backend único (Elysia/Node), banco MySQL (Prisma), "
        "quatro frontends e integrações externas."
    )
    doc.add_paragraph(
        "Diagrama (texto — compatível Mermaid):\n"
        "flowchart LR\n"
        "  subgraph clients [Clientes]\n"
        "    AC[App Cliente Expo]\n"
        "    CW[Client Web Next]\n"
        "  end\n"
        "  subgraph producers [Produtora]\n"
        "    AP[Produtor App Expo]\n"
        "    PW[Producer Web Next]\n"
        "  end\n"
        "  subgraph platform [Plataforma]\n"
        "    API[Backend Elysia]\n"
        "    DB[(MySQL)]\n"
        "    PF[pulse-face Python]\n"
        "  end\n"
        "  subgraph external [Externos]\n"
        "    PG[Pagar.me]\n"
        "    BR[Brevo e-mail]\n"
        "    BA[Better Auth]\n"
        "  end\n"
        "  AC --> API\n"
        "  CW --> API\n"
        "  AP --> API\n"
        "  PW --> API\n"
        "  API --> DB\n"
        "  API --> PG\n"
        "  API --> BR\n"
        "  API --> BA\n"
        "  API --> PF"
    )
    doc.add_paragraph(
        "Superfícies API: /api/client/v1 (B2C canônico), espelho legado na raiz; "
        "/api/producer/v1 (portal produtor + operação); /api/admin/v1 (Pulse Admin); "
        "/api/auth/* (Better Auth); /api/promoter (comissões)."
    )

    doc.add_heading("23. Recursos necessários para executar", level=1)
    doc.add_paragraph(
        "Infra mínima: processo Node (backend), MySQL, opcional pulse-face, "
        "deploys estáticos/SSR para Producer Web e Client Web, builds EAS para apps."
    )
    add_table_from_rows(
        doc,
        ["Categoria", "Exemplos (sem valores secretos)"],
        [
            ("Banco", "DATABASE_URL ou MYSQL_* — Prisma migrations"),
            ("Auth", "BETTER_AUTH_SECRET, BETTER_AUTH_URL"),
            ("HTTP", "PORT, CORS origins dos frontends"),
            ("Pagamentos", "PAGARME_SECRET_KEY, flags PAYMENTS_ENABLED nos apps"),
            ("E-mail", "BREVO_API_KEY, BREVO_SENDER_*, MAIL_PROVIDER, PRODUCER_WEB_URL"),
            ("Facial", "BIOMETRIC_*, FACIAL_*, PULSE_FACE_SERVICE_URL, PULSE_FACE_SERVICE_API_KEY"),
            ("KYC/Admin", "KYC_STORAGE_PATH, seed PULSE_ADMIN (bun run seed:pulse-admin)"),
            ("Jobs", "ENABLE_PAYOUT_RELEASE_JOB, ENABLE_FACE_GALLERY_PURGE_JOB"),
            ("Interno", "PULSE_INTERNAL_API_KEY, QR_SECRET"),
        ],
    )

    doc.add_heading("24. Catálogo de endpoints API (consolidado)", level=1)
    doc.add_paragraph(
        "Lista derivada dos controllers Elysia (backend/src). OpenAPI interativo: GET /swagger. "
        "Rotas legadas B2C na raiz espelham /api/client/v1 — preferir o prefixo canônico em novos clientes."
    )
    add_table_from_rows(
        doc,
        ["Domínio", "Método", "Caminho", "Descrição resumida"],
        API_ENDPOINT_ROWS,
    )


API_ENDPOINT_ROWS: list[list[str]] = [
    # Health
    ("Sistema", "GET", "/health", "Health check"),
    # Admin
    ("Admin v1", "POST", "/api/admin/v1/auth/login", "Login etapa 1 → OTP e-mail"),
    ("Admin v1", "POST", "/api/admin/v1/auth/login/verify-otp", "Login etapa 2 → token"),
    ("Admin v1", "POST", "/api/admin/v1/auth/logout", "Logout admin"),
    ("Admin v1", "GET", "/api/admin/v1/auth/me", "Sessão admin"),
    ("Admin v1", "GET", "/api/admin/v1/producers", "Listar produtoras + GMV"),
    ("Admin v1", "POST", "/api/admin/v1/producers", "Criar produtora (HU02)"),
    ("Admin v1", "POST", "/api/admin/v1/producers/:id/reset-password", "Reset senha produtor"),
    ("Admin v1", "GET", "/api/admin/v1/metrics/health", "Métricas checkout (HU03)"),
    ("Admin v1", "GET", "/api/admin/v1/payouts", "Listar repasses admin"),
    ("Admin v1", "GET", "/api/admin/v1/payouts/stats", "KPIs repasses"),
    ("Admin v1", "POST", "/api/admin/v1/payouts/events/:eventId/freeze", "Congelar repasse"),
    ("Admin v1", "POST", "/api/admin/v1/payouts/events/:eventId/unfreeze", "Descongelar"),
    ("Admin v1", "GET", "/api/admin/v1/refunds", "Listar estornos"),
    ("Admin v1", "GET", "/api/admin/v1/refunds/stats", "KPIs estornos"),
    ("Admin v1", "GET", "/api/admin/v1/refunds/producers/:producerId/events", "Eventos p/ estorno"),
    ("Admin v1", "GET", "/api/admin/v1/refunds/search-orders", "Buscar pedido"),
    ("Admin v1", "POST", "/api/admin/v1/refunds/validate", "Validar estorno"),
    ("Admin v1", "POST", "/api/admin/v1/refunds", "Processar estorno (HU05)"),
    ("Admin v1", "GET", "/api/admin/v1/compliance", "Documentos legais"),
    ("Admin v1", "POST", "/api/admin/v1/compliance/documents", "Publicar termo (HU06)"),
    ("Admin v1", "GET", "/api/admin/v1/kyc/queue", "Fila KYC"),
    ("Admin v1", "GET", "/api/admin/v1/kyc/documents/:id", "Detalhe KYC"),
    ("Admin v1", "GET", "/api/admin/v1/kyc/documents/:id/download", "Download documento"),
    ("Admin v1", "POST", "/api/admin/v1/kyc/documents/:id/approve", "Aprovar KYC"),
    ("Admin v1", "POST", "/api/admin/v1/kyc/documents/:id/reject", "Rejeitar KYC"),
    # Producer
    ("Producer v1", "POST", "/api/producer/v1/auth/login", "Login produtor"),
    ("Producer v1", "POST", "/api/producer/v1/auth/onboarding/*", "Fluxo onboarding/OTP"),
    ("Producer v1", "GET|PATCH", "/api/producer/v1/profile/*", "Perfil produtor"),
    ("Producer v1", "GET|POST|PATCH|DELETE", "/api/producer/v1/team/*", "Equipe e convites"),
    ("Producer v1", "GET|POST|PATCH", "/api/producer/v1/events/*", "CRUD eventos"),
    ("Producer v1", "GET|POST|PATCH", "/api/producer/v1/events/:id/commercial/*", "Setores e lotes"),
    ("Producer v1", "GET|POST", "/api/producer/v1/finance/*", "Financeiro e saques"),
    ("Producer v1", "GET|POST", "/api/producer/v1/kyc/documents/*", "KYC titular upload"),
    ("Producer v1", "GET|POST", "/api/producer/v1/operation/*", "Check-in QR/facial/lista"),
    ("Producer v1", "GET|POST", "/api/producer/v1/compliance/*", "Aceite termos produtor"),
    # Client v1 (canônico — espelho na raiz)
    ("Client v1", "POST", "/api/client/v1/auth/*", "Cadastro/login B2C"),
    ("Client v1", "GET", "/api/client/v1/events/*", "Catálogo e detalhe"),
    ("Client v1", "GET|POST", "/api/client/v1/tickets/*", "Carteira, transfer, cancel"),
    ("Client v1", "GET|POST", "/api/client/v1/checkout/*", "Sessão de compra"),
    ("Client v1", "GET|POST|DELETE", "/api/client/v1/payment/*", "Cartões e Pix/cartão"),
    ("Client v1", "GET|POST|DELETE", "/api/client/v1/biometry/*", "Facial enrollment"),
    ("Client v1", "GET|POST", "/api/client/v1/compliance/*", "Termos B2C"),
    ("Promoter", "GET", "/api/promoter/*", "Vendas e comissões promoter"),
    ("Auth", "ALL", "/api/auth/*", "Better Auth handler"),
    ("Interno", "POST", "/internal/facial-*", "Jobs galeria/retenção (API key)"),
]


def build_client_module_table(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("20. Módulos B2C — mapa de implementação", level=1)
    doc.add_paragraph(
        "Consolidado para decisão bug vs comportamento esperado. "
        "Referência: docs/product/especificacao_funcional_mvp_ingressos.docx."
    )
    add_table_from_rows(
        doc,
        ["Módulo", "Regra-chave", "Produtor App", "Producer Web", "Client App", "Client Web"],
        [
            ("Eventos (vitrine)", "Status e janela de venda", "—", "—", "[IMPLEMENTADO]", "[IMPLEMENTADO]"),
            ("Setores e lotes", "Máx. 4 ingressos/evento/CPF", "—", "[IMPLEMENTADO]", "[IMPLEMENTADO]", "[IMPLEMENTADO] seleção"),
            ("Checkout", "Reserva 10 min, 3 tentativas", "—", "—", "[PARCIAL] flag pagamento", "[PENDENTE]"),
            ("Pagamentos", "Pix 5% desconto taxa; cartão 4x", "—", "—", "[PARCIAL]", "[PENDENTE]"),
            ("Ingressos / carteira", "Só após PAID", "—", "—", "[IMPLEMENTADO]", "[PENDENTE]"),
            ("Facial", "Principal; QR fallback", "[IMPLEMENTADO] Access", "[PENDENTE]", "[IMPLEMENTADO]", "[PENDENTE]"),
            ("Cancelamento", "24h antes; sem pós-USD", "—", "—", "[IMPLEMENTADO]", "[PENDENTE]"),
            ("Financeiro produtor", "Retenção D+1 pós término*", "—", "[PARCIAL]", "—", "—"),
            ("VIP assinatura", "Plano recorrente", "[PENDENTE]", "[PARCIAL] stub", "[PARCIAL] mock", "—"),
            ("Promoter", "Comissões no app cliente", "—", "[IMPLEMENTADO] convite", "[IMPLEMENTADO]", "—"),
        ],
    )
    doc.add_paragraph(
        "*Gatilho de liberação: job ReleaseRetainedPayoutsUseCase — 24h após endDate do evento "
        "(ver product/payout-policies.md — canônico D+1)."
    )


def patch_hu91(doc: Document) -> None:
    for p in doc.paragraphs:
        if HU91_STRIKE in p.text:
            strike_paragraph(p)
            p.add_run(f" {HU91_REPLACE}")
        if HU91_ADMIN_STRIKE in p.text:
            replace_in_paragraph(p, {HU91_ADMIN_STRIKE: HU91_ADMIN_REPLACE})


def renumber_old_post_mvp(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("14. Evoluções pós-MVP"):
            p.text = "18. Evoluções pós-MVP (substituído — ver seção 18 atualizada)"
            strike_paragraph(p)
        elif t.startswith("15. Observações funcionais"):
            p.text = "19. Observações funcionais e técnicas sobre o módulo financeiro"


def main() -> None:
    if not SRC_PRODUCER.exists():
        raise SystemExit(f"Fonte não encontrada: {SRC_PRODUCER}")

    shutil.copy2(SRC_PRODUCER, OUT)
    doc = Document(OUT)

    # Capa
    for p in doc.paragraphs[:8]:
        if "Especificação Funcional Unificada" in p.text or "Pulse!" in p.text[:30]:
            if "Unificada" in p.text:
                p.text = "Especificação Funcional Unificada — Ecossistema Pulse"
        if "MVP — App do Produtor" in p.text:
            p.text = (
                f"MVP — Produtor App · Producer Web · App Cliente · Client Web "
                f"(atualização {TODAY.strftime('%d/%m/%Y')})"
            )
        if "Documento base para definição" in p.text or "Documento único" in p.text:
            p.text = (
                "Documento único de regras de negócio, fluxos, permissões e backlog. "
                "Referências: docs/RBAC.md, docs/CHECKOUT_COMPLIANCE.md, "
                "product/global-business-rules.md, product/payout-policies.md, código (apps + backend)."
            )

    # Glossário após primeiro parágrafo substantivo
    intro_idx = find_paragraph(doc, "Documento único")
    if intro_idx is not None:
        doc.paragraphs[intro_idx].insert_paragraph_before(GLOSSARIO)
        doc.paragraphs[intro_idx].insert_paragraph_before(LEGENDA)

    # Substituições globais
    for p in doc.paragraphs:
        t = p.text
        for old, new in STATUS_REPLACEMENTS.items():
            if old in t:
                replace_in_paragraph(p, {old: new})
        for emoji, tag in {
            "✅": "[IMPLEMENTADO]",
            "🟡": "[PARCIAL]",
            "⏳": "[PENDENTE]",
            "🔵": "[PENDENTE]",
        }.items():
            if emoji in t:
                replace_in_paragraph(p, {emoji: tag})

    if doc.tables:
        t0 = doc.tables[0]
        if len(t0.rows) >= 4:
            t0.rows[1].cells[1].text = (
                "Produtor App + Producer Web + App Cliente + Client Web + backend"
            )
            t0.rows[2].cells[1].text = f"2.1 — de/para código e docs ({TODAY.strftime('%Y-%m')})"
            t0.rows[3].cells[1].text = (
                "product/rbac.md, global-business-rules.md, checkout-compliance.md"
            )
        if len(doc.tables) > 1:
            patch_roadmap_table(doc.tables[1])

    for idx, patch in FEATURE_TABLE_PATCHES.items():
        if idx < len(doc.tables):
            add_platform_columns_to_feature_table(doc.tables[idx], patch)

    for table in doc.tables:
        for row in table.rows:
            line = " | ".join(c.text for c in row.cells)
            for frag in STRIKE_CONTAINS:
                if frag in line or frag.replace("⏳", "[PENDENTE]") in line:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            strike_paragraph(para)

    for p in doc.paragraphs:
        if p.text.strip() == "Status: [PENDENTE] Não implementado":
            strike_paragraph(p)

    patch_hu91(doc)
    renumber_old_post_mvp(doc)

    # Remove inserções anteriores mal ordenadas (14–20)
    remove_section_range(doc, "17. Legado", "18. Evoluções")
    remove_section_range(doc, "14. Ecossistema", "17. Legado")
    remove_section_range(doc, "20. Módulos B2C", "20. Módulos B2C")  # noop if only title

    # Remove §20 table block if exists (heading + table is harder — remove heading and following paras until §18/19)
    idx20 = find_paragraph(doc, "20. Módulos B2C")
    if idx20 is not None:
        end = find_paragraph(doc, "18. Evoluções") or find_paragraph(doc, "19. Observações")
        if end and end > idx20:
            for p in list(doc.paragraphs[idx20:end]):
                delete_paragraph(p)

    # Anchor: §18 ou antigo 14 Evoluções
    anchor_idx = find_paragraph(doc, "18. Evoluções")
    if anchor_idx is None:
        anchor_idx = find_paragraph(doc, "14. Evoluções pós-MVP")
    if anchor_idx is None:
        anchor_idx = len(doc.paragraphs) - 1

    anchor = doc.paragraphs[anchor_idx]

    # Remove corpo antigo do §18 (bullets obsoletos) até §19
    idx19 = find_paragraph(doc, "19. Observações")
    if idx19 and idx19 > anchor_idx:
        for p in list(doc.paragraphs[anchor_idx + 1 : idx19]):
            txt = p.text.strip()
            if txt and not txt.startswith("19."):
                delete_paragraph(p)
        anchor = doc.paragraphs[find_paragraph(doc, "18. Evoluções") or anchor_idx]

    # Renomear heading §18 antigo se ainda existir
    if anchor.text.strip().startswith("14. Evoluções"):
        anchor.text = "18. Evoluções pós-MVP (legado — ver abaixo)"
        strike_paragraph(anchor)

    # Inserir §14–17 e §18 atualizado antes do anchor legado
    insert_blocks_before(anchor, ecosystem_blocks())
    insert_blocks_before(anchor, post_mvp_blocks())

    # Remover heading §18 legado duplicado (strikethrough)
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("18. Evoluções pós-MVP (substituído") or t == "18. Evoluções pós-MVP":
            if p.style and p.style.name.startswith("Heading"):
                delete_paragraph(p)

    # §20 tabela B2C + apêndices 22–24
    build_client_module_table(doc)
    appendix_architecture(doc)

    # §21 Pulse Admin entre §20 e §22
    idx22 = find_paragraph(doc, "22. Arquitetura")
    if idx22 is not None:
        insert_blocks_before(doc.paragraphs[idx22], pulse_admin_blocks())
        idx217 = find_paragraph(doc, "21.7 Mapa HU")
        if idx217 is not None:
            anchor_p = doc.paragraphs[idx217]
            p88 = anchor_p.insert_paragraph_before("21.8 Tabela HU × rota × status")
            try:
                p88.style = "Heading 2"
            except Exception:
                pass
            insert_table_after_paragraph(
                doc,
                anchor_p,
                ["HU", "Fluxo", "Rota/UI", "Status"],
                pulse_admin_hu_table_rows(),
            )

    doc.save(OUT)
    print(f"Saved {OUT}")


if __name__ == "__main__":
    main()
